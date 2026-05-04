import re
import sys
import zipfile
from pathlib import Path

from lxml import etree
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT


NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
}

RED = RGBColor(192, 0, 0)


def tag_name(element):
    return etree.QName(element).localname


def get_text(element):
    texts = element.xpath(".//w:t/text() | .//w:delText/text()", namespaces=NS)
    return "".join(texts)


def paragraph_text(paragraph):
    return get_text(paragraph).strip()


def is_red_run(run):
    colors = run.xpath(".//w:rPr/w:color/@w:val", namespaces=NS)
    for color in colors:
        color = color.lower()
        if color in ["ff0000", "c00000", "c0392b", "e60000", "b00000", "a00000", "red"]:
            return True
    return False


def is_strike_run(run):
    return bool(run.xpath(".//w:rPr/w:strike | .//w:rPr/w:dstrike", namespaces=NS))


def is_underline_run(run):
    return bool(run.xpath(".//w:rPr/w:u", namespaces=NS))


def collect_runs_from_element(element, inside_change=None):
    results = []
    local = tag_name(element)

    if local == "del":
        inside_change = "delete"
    elif local == "ins":
        inside_change = "insert"

    if local == "r":
        text = get_text(element)

        if text:
            visual_red = is_red_run(element)
            visual_strike = is_strike_run(element)
            visual_underline = is_underline_run(element)

            if inside_change == "delete":
                results.append({
                    "text": text,
                    "red": True,
                    "strike": True,
                    "underline": False,
                    "is_redline": True,
                })
            elif inside_change == "insert":
                results.append({
                    "text": text,
                    "red": True,
                    "strike": False,
                    "underline": True,
                    "is_redline": True,
                })
            elif visual_red or visual_strike:
                results.append({
                    "text": text,
                    "red": True,
                    "strike": visual_strike,
                    "underline": visual_underline,
                    "is_redline": True,
                })
            else:
                results.append({
                    "text": text,
                    "red": False,
                    "strike": False,
                    "underline": False,
                    "is_redline": False,
                })

        return results

    for child in element:
        results.extend(collect_runs_from_element(child, inside_change))

    return results


def paragraph_has_redline(paragraph_info):
    return any(run["is_redline"] for run in paragraph_info["runs"])


def detect_page(text, current_page):
    match = re.search(r"Page Number\s*:\s*(\d+)", text, re.IGNORECASE)
    if match:
        return int(match.group(1))
    return current_page


def detect_lesson(text, current_lesson):
    clean = text.strip()
    lowered = clean.lower()

    # Avoid incorrectly using "Lesson Type :" as the lesson name.
    if lowered.startswith("lesson type"):
        return current_lesson

    if re.match(r"^lesson\s+\d+\s*:", clean, re.IGNORECASE):
        return clean

    if lowered.startswith("lesson name:"):
        return clean.split(":", 1)[1].strip()

    return current_lesson


def detect_template(text, current_template):
    clean = text.strip()
    lowered = clean.lower()

    known_templates = [
        "click and reveal template",
        "binary list template",
        "question and answer template",
        "text and image template",
        "video template",
        "hotspot image template",
    ]

    for template in known_templates:
        if lowered.startswith(template):
            return clean

    if lowered.startswith("select type"):
        return clean.replace("Select Type", "").replace(":", "").strip() or current_template

    if lowered.startswith("template used:"):
        return clean.split(":", 1)[1].strip()

    if lowered.startswith("template type:"):
        return clean.split(":", 1)[1].strip()

    if lowered.startswith("template:"):
        return clean.split(":", 1)[1].strip()

    return current_template


def is_question_start(text):
    clean = text.strip()
    return bool(
        re.match(r"^\d+\.\s*Question\s*:", clean, re.IGNORECASE)
        or re.match(r"^Question\s*:", clean, re.IGNORECASE)
    )


def is_page_header(text):
    return bool(re.match(r"^Page Number\s*:", text.strip(), re.IGNORECASE))


def is_lesson_header(text):
    clean = text.strip()
    return bool(re.match(r"^Lesson\s+\d+\s*:", clean, re.IGNORECASE))


def is_template_header(text):
    clean = text.strip().lower()

    known_templates = [
        "click and reveal template",
        "binary list template",
        "question and answer template",
        "text and image template",
        "video template",
        "hotspot image template",
    ]

    return any(clean.startswith(template) for template in known_templates)


def is_quiz_end_boundary(text):
    clean = text.strip().lower()

    endings = [
        "number of questions needed to pass",
        "passed message",
        "failed message",
        "feedback for correct",
        "feedback for incorrect",
        "feedback for partial",
        "reveal content for desktop",
        "reveal content for mobile",
        "audio transcript text",
    ]

    return any(clean.startswith(item) for item in endings)


def is_major_boundary(text):
    if not text.strip():
        return False

    return (
        is_page_header(text)
        or is_lesson_header(text)
        or is_template_header(text)
        or is_quiz_end_boundary(text)
    )


def add_docx_run(paragraph, text, red=False, strike=False, underline=False, bold=False):
    if not text:
        return

    run = paragraph.add_run(text)
    run.font.size = Pt(14)

    if red:
        run.font.color.rgb = RED

    if strike:
        run.font.strike = True

    if underline:
        run.font.underline = True

    if bold:
        run.bold = True


def add_paragraph_runs(cell, paragraph_info, redline_only=False):
    p = cell.add_paragraph()

    for run in paragraph_info["runs"]:
        if redline_only and not run["is_redline"]:
            continue

        add_docx_run(
            p,
            run["text"],
            red=run["red"],
            strike=run["strike"],
            underline=run["underline"],
        )


def block_has_redline(block):
    return any(paragraph_has_redline(p) for p in block["paragraphs"])


def create_block(page, lesson, template, block_type="Redline Block"):
    return {
        "page": page,
        "lesson": lesson,
        "template": template,
        "block_type": block_type,
        "paragraphs": [],
    }


def should_continue_redline_group(text):
    clean = text.strip()

    if not clean:
        return True

    if is_major_boundary(clean):
        return False

    if is_question_start(clean):
        return False

    return True


def extract_redline_blocks(docx_path):
    results = []

    current_page = 1
    current_lesson = "Unknown"
    current_template = "Unknown"

    active_question_block = None
    active_redline_group = None
    quiet_gap_count = 0

    with zipfile.ZipFile(docx_path) as docx:
        xml = docx.read("word/document.xml")

    root = etree.fromstring(xml)
    paragraphs = root.xpath(".//w:p", namespaces=NS)

    for paragraph in paragraphs:
        text = paragraph_text(paragraph)
        runs = collect_runs_from_element(paragraph)

        paragraph_info = {
            "text": text,
            "runs": runs,
        }

        has_redline = paragraph_has_redline(paragraph_info)

        # Always update metadata from document text.
        if text:
            current_page = detect_page(text, current_page)
            current_lesson = detect_lesson(text, current_lesson)
            current_template = detect_template(text, current_template)

        # Close normal redline group when we reach a major boundary.
        if active_redline_group and text and is_major_boundary(text):
            if block_has_redline(active_redline_group):
                results.append(active_redline_group)
            active_redline_group = None
            quiet_gap_count = 0

        # If a new question starts, close any existing group first.
        if text and is_question_start(text):
            if active_redline_group and block_has_redline(active_redline_group):
                results.append(active_redline_group)
            active_redline_group = None
            quiet_gap_count = 0

            if active_question_block and block_has_redline(active_question_block):
                results.append(active_question_block)

            active_question_block = create_block(
                current_page,
                current_lesson,
                current_template,
                "Question Block"
            )
            active_question_block["paragraphs"].append(paragraph_info)
            continue

        # If inside a question block, keep collecting until a major boundary.
        if active_question_block:
            if text and is_major_boundary(text):
                if block_has_redline(active_question_block):
                    results.append(active_question_block)
                active_question_block = None
            else:
                active_question_block["paragraphs"].append(paragraph_info)
                continue

        # Group related non-question redline paragraphs.
        if has_redline:
            if not active_redline_group:
                active_redline_group = create_block(
                    current_page,
                    current_lesson,
                    current_template,
                    "Redline Group"
                )

            active_redline_group["paragraphs"].append(paragraph_info)
            quiet_gap_count = 0
            continue

        # If we are inside a redline group, keep nearby context lines.
        # This helps combine items inside the same visual box/list.
        if active_redline_group:
            if should_continue_redline_group(text) and quiet_gap_count < 2:
                active_redline_group["paragraphs"].append(paragraph_info)

                if text.strip():
                    quiet_gap_count += 1

                continue

            if block_has_redline(active_redline_group):
                results.append(active_redline_group)

            active_redline_group = None
            quiet_gap_count = 0

    # Close final active blocks.
    if active_question_block and block_has_redline(active_question_block):
        results.append(active_question_block)

    if active_redline_group and block_has_redline(active_redline_group):
        results.append(active_redline_group)

    return results


def set_cell_header(cell, text):
    p = cell.paragraphs[0]
    add_docx_run(p, text, bold=True)


def remove_empty_first_paragraph(cell):
    if cell.paragraphs and not cell.paragraphs[0].text:
        p = cell.paragraphs[0]._element
        p.getparent().remove(p)


def create_meta_box(doc, page, lesson, template):
    meta_table = doc.add_table(rows=3, cols=1)
    meta_table.style = "Table Grid"
    meta_table.alignment = WD_TABLE_ALIGNMENT.LEFT

    cells = [
        meta_table.cell(0, 0),
        meta_table.cell(1, 0),
        meta_table.cell(2, 0),
    ]

    labels = [
        f"Page Number: {page}",
        lesson,
        template,
    ]

    for cell, label in zip(cells, labels):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        paragraph = cell.paragraphs[0]
        paragraph.alignment = 0
        run = paragraph.add_run(label)
        run.bold = True
        run.font.size = Pt(10)

    return meta_table


def create_two_column_docx(results, output_file):
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    title = doc.add_paragraph()
    add_docx_run(title, "Redline Review Output", bold=True)

    intro = doc.add_paragraph()
    add_docx_run(
        intro,
        "Each table groups related redline changes together. For questions, the full question and answer set appears in one block."
    )

    if not results:
        doc.add_paragraph("No redline changes found.")
        doc.save(output_file)
        return

    for item in results:
        # Metadata box on the LEFT, similar to the source document.
        create_meta_box(
            doc,
            item["page"],
            item["lesson"],
            item["template"]
        )

        doc.add_paragraph("")

        table = doc.add_table(rows=2, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

        set_cell_header(table.cell(0, 0), "Original Content")
        set_cell_header(table.cell(0, 1), "Redline Content Only")

        original_cell = table.cell(1, 0)
        redline_cell = table.cell(1, 1)

        for paragraph_info in item["paragraphs"]:
            add_paragraph_runs(original_cell, paragraph_info, redline_only=False)

        for paragraph_info in item["paragraphs"]:
            if paragraph_has_redline(paragraph_info):
                add_paragraph_runs(redline_cell, paragraph_info, redline_only=True)

        remove_empty_first_paragraph(original_cell)
        remove_empty_first_paragraph(redline_cell)

        doc.add_paragraph("")

    doc.save(output_file)


def main():
    if len(sys.argv) != 3:
        print("Usage: python scripts/extract_redlines.py <input_folder> <output_docx>")
        sys.exit(1)

    input_folder = Path(sys.argv[1])
    output_file = Path(sys.argv[2])
    output_file.parent.mkdir(parents=True, exist_ok=True)

    docx_files = sorted(input_folder.glob("**/*.docx"))

    if not docx_files:
        print("ERROR: No .docx files found in the input folder.")
        sys.exit(2)

    all_results = []

    for docx_file in docx_files:
        print(f"Processing: {docx_file}")
        all_results.extend(extract_redline_blocks(docx_file))

    create_two_column_docx(all_results, output_file)

    print(f"Created report: {output_file}")
    print(f"Redline blocks found: {len(all_results)}")


if __name__ == "__main__":
    main()
